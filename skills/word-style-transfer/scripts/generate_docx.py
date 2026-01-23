import argparse
import sys
import os
from docx import Document
from docx.shared import Pt

def clean_document_content(doc):
    """
    Clears the content of the document but preserves styles.
    This is a bit tricky in python-docx as there isn't a single 'clear' method.
    We iterate and remove body elements.
    """
    # This is a naive approach: we can't easily "empty" a docx while keeping all headers/footers/styles perfectly 
    # using just python-docx without diving into XML. 
    # A safer approach for "Style Transfer" using python-docx is:
    # 1. Open source doc.
    # 2. Delete all paragraphs in body.
    # 3. Delete all tables in body.
    
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
        
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)
        t._t = t._element = None

def parse_markdown_to_docx(doc, markdown_path):
    """
    Reads markdown file and adds paragraphs to doc using styles.
    """
    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[2:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[2:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line[0].isdigit() and '. ' in line[:4]:
            # Simple check for numbered lists "1. "
            parts = line.split('. ', 1)
            if len(parts) > 1:
                doc.add_paragraph(parts[1], style='List Number')
            else:
                doc.add_paragraph(line)
        else:
            # Normal text
            doc.add_paragraph(line)

def main():
    parser = argparse.ArgumentParser(description='Generate a DOCX from Markdown using a template for styles.')
    parser.add_argument('--template', required=True, help='Path to the source DOCX to use as a style template.')
    parser.add_argument('--content', required=True, help='Path to the content Markdown file.')
    parser.add_argument('--output', required=True, help='Path to save the generated DOCX.')

    args = parser.parse_args()

    if not os.path.exists(args.template):
        print(f"Error: Template file not found: {args.template}")
        sys.exit(1)

    if not os.path.exists(args.content):
        print(f"Error: Content file not found: {args.content}")
        sys.exit(1)

    try:
        # Load the template
        doc = Document(args.template)
        
        # Clear existing content (to use it as a blank canvas with styles)
        clean_document_content(doc)
        
        # Add new content
        parse_markdown_to_docx(doc, args.content)
        
        # Save
        doc.save(args.output)
        print(f"Successfully generated: {args.output}")

    except Exception as e:
        print(f"An error occurred: {str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()
